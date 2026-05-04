from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document
from pypdf import PdfReader

MAX_UPLOAD_BYTES = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".txt", ".docx", ".pdf"}


class DocumentExtractionError(Exception):
    pass


@dataclass
class ExtractionResult:
    text: str
    detected_type: str
    warnings: list[str]


def extract_uploaded_document_text(filename: str, content: bytes) -> ExtractionResult:
    suffix = Path(filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise DocumentExtractionError("Unsupported file type. Allowed: .txt, .docx, .pdf")
    if not content:
        raise DocumentExtractionError("Uploaded file is empty")
    if len(content) > MAX_UPLOAD_BYTES:
        raise DocumentExtractionError("File is too large (max 10 MB)")

    if suffix == ".txt":
        try:
            text = _extract_txt(content)
        except Exception as exc:
            raise DocumentExtractionError(f"Could not read TXT file: {exc}") from exc
        return ExtractionResult(text=text, detected_type="txt", warnings=[] if text else ["No text found in TXT file"])

    if suffix == ".docx":
        try:
            text = _extract_docx(content)
        except Exception as exc:
            raise DocumentExtractionError(f"Could not read DOCX file: {exc}") from exc
        return ExtractionResult(text=text, detected_type="docx", warnings=[] if text else ["No text found in DOCX file"])

    try:
        return _extract_pdf_with_docling_fallback(content)
    except Exception as exc:
        raise DocumentExtractionError(f"Could not read PDF file: {exc}") from exc


def _extract_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8").strip()
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="ignore").strip()


def _extract_docx(content: bytes) -> str:
    doc = Document(BytesIO(content))
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        value = paragraph.text.strip()
        if value:
            lines.append(value)
    for table in doc.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_values:
                lines.append(" | ".join(row_values))
    return "\n".join(lines).strip()


def _extract_pdf_with_docling_fallback(content: bytes) -> ExtractionResult:
    warnings: list[str] = []

    docling_text = _extract_pdf_with_docling(content)
    if docling_text:
        return ExtractionResult(text=docling_text, detected_type="pdf", warnings=warnings)

    warnings.append("Docling could not extract PDF text; used pypdf fallback")
    fallback_text = _extract_pdf_with_pypdf(content)
    if fallback_text:
        return ExtractionResult(text=fallback_text, detected_type="pdf", warnings=warnings)

    warnings.append("No readable text found in PDF (possibly scanned image PDF)")
    return ExtractionResult(text="", detected_type="pdf", warnings=warnings)


def _extract_pdf_with_docling(content: bytes) -> str:
    try:
        from docling.document_converter import DocumentConverter
    except Exception:
        return ""

    try:
        with NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
            tmp.write(content)
            tmp.flush()
            converter = DocumentConverter()
            result = converter.convert(tmp.name)
            document = getattr(result, "document", None)
            if document is None:
                return ""
            markdown = document.export_to_markdown()
            return markdown.strip() if isinstance(markdown, str) else ""
    except Exception:
        return ""


def _extract_pdf_with_pypdf(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
    except Exception:
        return ""

    chunks: list[str] = []
    for page in reader.pages:
        try:
            value = page.extract_text() or ""
        except Exception:
            value = ""
        value = value.strip()
        if value:
            chunks.append(value)
    return "\n\n".join(chunks).strip()
