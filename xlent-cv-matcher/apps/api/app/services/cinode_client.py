from __future__ import annotations

from io import BytesIO
from typing import Any

import httpx
from docx import Document

from app.core.config import Settings
from app.services.cinode_directory import _resolve_api_authorization, normalize_auth_value


class CinodePublishError(Exception):
    pass


def _safe_text(value: Any, limit: int = 4000) -> str:
    text = str(value or "").replace("\r", " ").strip()
    return text[:limit]


def _extract_company_and_user(payload: dict[str, Any]) -> tuple[str | None, str | None]:
    sources: list[dict[str, Any]] = []
    for key in ["source_resume", "source_profile"]:
        value = payload.get(key)
        if isinstance(value, dict):
            sources.append(value)
    sources.append(payload)

    company_id: str | None = None
    user_id: str | None = None

    for source in sources:
        if company_id is None:
            value = source.get("companyId") or source.get("company_id")
            if value is not None and str(value).strip():
                company_id = str(value).strip()

        if user_id is None:
            value = source.get("companyUserId") or source.get("company_user_id")
            if value is not None and str(value).strip():
                user_id = str(value).strip()

        if user_id is None:
            value = source.get("id")
            if value is not None and str(value).strip():
                user_id = str(value).strip()

    return company_id, user_id


def _resolve_publish_path(path_template: str, company_id: str, user_id: str) -> str:
    path = (path_template or "").strip() or "/v0.1/companies/{companyId}/users/{companyUserId}/profile/import"
    path = path.replace("{companyId}", company_id).replace("{companyUserId}", user_id)
    path = path.replace("{companyid}", company_id).replace("{companyuserid}", user_id)
    return path


def _extract_cv_paragraphs(payload: dict[str, Any], limit: int = 16) -> list[str]:
    paragraphs: list[str] = []
    source_resume = payload.get("source_resume")
    if isinstance(source_resume, dict):
        resume = source_resume.get("resume")
        if isinstance(resume, dict):
            blocks = resume.get("blocks")
            if isinstance(blocks, list):
                for block in blocks[:40]:
                    if not isinstance(block, dict):
                        continue
                    title = _safe_text(block.get("title"), limit=160)
                    description = _safe_text(block.get("description"), limit=800)
                    if title and description:
                        paragraphs.append(f"{title}: {description}")
                    elif description:
                        paragraphs.append(description)
                    elif title:
                        paragraphs.append(title)
                    if len(paragraphs) >= limit:
                        break
    return paragraphs


def _build_docx_from_payload(title: str, payload: dict[str, Any]) -> bytes:
    document = Document()
    document.add_heading(_safe_text(title, limit=200) or "CV", level=0)

    name = _safe_text(payload.get("name"), limit=200)
    role = _safe_text(payload.get("title"), limit=200)
    location = _safe_text(payload.get("location"), limit=200)
    if name:
        document.add_paragraph(f"Navn: {name}")
    if role:
        document.add_paragraph(f"Tittel: {role}")
    if location:
        document.add_paragraph(f"Lokasjon: {location}")

    summary = _safe_text(payload.get("summary"), limit=4000)
    if summary:
        document.add_heading("Sammendrag", level=1)
        document.add_paragraph(summary)

    skills = payload.get("skills")
    if isinstance(skills, list) and skills:
        document.add_heading("Kompetanse", level=1)
        for skill in skills[:60]:
            text = _safe_text(skill, limit=160)
            if text:
                document.add_paragraph(text, style="List Bullet")

    paragraphs = _extract_cv_paragraphs(payload)
    if paragraphs:
        document.add_heading("Erfaring og prosjekter", level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    tailored = payload.get("tailored_sections")
    if isinstance(tailored, list):
        tailored_rows: list[tuple[str, str]] = []
        for entry in tailored[:30]:
            if not isinstance(entry, dict):
                continue
            section = _safe_text(entry.get("section_type"), limit=80) or "other"
            text = _safe_text(entry.get("suggested_text"), limit=3000)
            if not text:
                continue
            tailored_rows.append((section, text))
        if tailored_rows:
            document.add_heading("Tilpassede tekstforslag", level=1)
            for section, text in tailored_rows:
                document.add_heading(section.capitalize(), level=2)
                document.add_paragraph(text)

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_docx_from_payload(title: str, payload: dict[str, Any]) -> bytes:
    return _build_docx_from_payload(title=title, payload=payload)


def _normalize_location_header(location_header: str, base_url: str) -> str:
    if location_header.lower().startswith("http://") or location_header.lower().startswith("https://"):
        return location_header
    return f"{base_url.rstrip('/')}/{location_header.lstrip('/')}"


def _interpret_operation_status(payload: Any) -> tuple[bool, bool, str]:
    if isinstance(payload, str):
        text = payload.strip()
        lowered = text.lower()
        if any(token in lowered for token in ["failed", "error"]):
            return True, False, text
        if any(token in lowered for token in ["completed", "done", "success"]):
            return True, True, text
        return False, False, text or "Operation in progress"

    if isinstance(payload, dict):
        lowered_keys = {str(k).lower(): v for k, v in payload.items()}
        status_text = ""
        for key in ["status", "state", "operationstatus", "message", "description"]:
            value = lowered_keys.get(key)
            if isinstance(value, str) and value.strip():
                status_text = value.strip()
                break

        completed = False
        for key in ["iscompleted", "completed", "done", "finished"]:
            value = lowered_keys.get(key)
            if isinstance(value, bool):
                completed = value
                break

        success_value: bool | None = None
        for key in ["issuccess", "success", "succeeded"]:
            value = lowered_keys.get(key)
            if isinstance(value, bool):
                success_value = value
                break

        status_lower = status_text.lower()
        if any(token in status_lower for token in ["failed", "error"]):
            return True, False, status_text or "Operation failed"
        if any(token in status_lower for token in ["completed", "done", "success"]):
            return True, True, status_text or "Operation completed"

        if completed:
            return True, bool(success_value is not False), status_text or "Operation completed"

        if success_value is True and status_text:
            return True, True, status_text

        return False, False, status_text or "Operation in progress"

    return False, False, "Operation in progress"


def _poll_profile_import_status(base_url: str, auth_header: str, operation_url: str) -> tuple[str, str]:
    url = _normalize_location_header(operation_url, base_url)
    headers = {"Authorization": auth_header}

    with httpx.Client(timeout=20) as client:
        for attempt in range(30):
            if attempt > 0:
                import time

                time.sleep(3.0)

            response = client.get(url, headers=headers)
            raw_text = (response.text or "").strip()

            if response.status_code >= 400:
                raise CinodePublishError(
                    f"Import status check failed at {url}: HTTP {response.status_code}. Response: {raw_text[:500] or '<empty>'}"
                )

            parsed: Any
            try:
                parsed = response.json()
            except Exception:
                parsed = raw_text

            completed, success, detail = _interpret_operation_status(parsed)
            if completed:
                if success:
                    return "success", detail
                return "failed", detail

    return "in_progress", "Import operation is still running"


def publish_to_cinode(
    settings: Settings,
    payload: dict[str, Any],
    title: str,
    dry_run: bool,
    base_url_override: str | None = None,
    auth_value_override: str | None = None,
) -> dict[str, Any]:
    base_url = base_url_override or settings.cinode_base_url
    auth_value = auth_value_override or settings.cinode_api_token

    if dry_run:
        return {
            "published": False,
            "dry_run": True,
            "target_url": base_url,
            "external_id": None,
            "detail": "Dry run: payload validated, no request sent to Cinode.",
        }

    if not settings.enable_cinode_publish:
        raise CinodePublishError("Cinode publish is disabled by configuration")

    if not base_url or not auth_value:
        raise CinodePublishError("Missing CINODE_BASE_URL or CINODE_API_TOKEN")

    company_id, user_id = _extract_company_and_user(payload)
    if not company_id or not user_id:
        raise CinodePublishError(
            "Could not resolve company/user id from payload. Expected source_profile/source_resume with companyId and companyUserId."
        )

    publish_path = _resolve_publish_path(settings.cinode_publish_path, company_id, user_id)
    url = f"{base_url.rstrip('/')}/{publish_path.lstrip('/')}"
    try:
        api_auth, _ = _resolve_api_authorization(base_url, auth_value)
    except Exception as exc:
        raise CinodePublishError(f"Credential bootstrap failed: {exc}") from exc

    docx_bytes = _build_docx_from_payload(title=title, payload=payload)
    filename = f"{_safe_text(title, 80).replace(' ', '_') or 'cv'}.docx"
    files = {
        "File": (
            filename,
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    form_data = {
        "ImportSkills": "true",
        "MapSkillExperienceYearsToLevel": "false",
    }
    headers = {
        "Authorization": api_auth,
    }

    with httpx.Client(timeout=30, follow_redirects=True) as client:
        response = client.post(url, headers=headers, files=files, data=form_data)

    if response.status_code >= 400:
        response_text = (response.text or "").strip()
        if not response_text:
            response_text = "<empty response body>"
        raise CinodePublishError(
            f"Cinode publish failed at {url}: HTTP {response.status_code}. Response: {response_text[:500]}"
        )

    operation_location = response.headers.get("Location") or response.headers.get("location")
    detail = "Import operation started"
    published = False
    external_id = None

    if response.status_code == 202 and operation_location:
        operation_state, operation_detail = _poll_profile_import_status(base_url, headers["Authorization"], operation_location)
        external_id = operation_location
        if operation_state == "success":
            published = True
            detail = f"Import completed. {operation_detail}"
        elif operation_state == "failed":
            raise CinodePublishError(
                "Cinode import operation failed. "
                f"Operation status endpoint: {operation_location}. Detail: {operation_detail}"
            )
        else:
            published = False
            detail = (
                "Import accepted by Cinode and is still processing. "
                f"Operation status endpoint: {operation_location}. Detail: {operation_detail}"
            )
    elif response.status_code in {200, 201, 204}:
        published = True
        detail = "Import completed"
    else:
        # Any non-error response still means request accepted by Cinode.
        published = response.status_code < 400
        detail = f"Import accepted with status {response.status_code}"

    return {
        "published": published,
        "dry_run": False,
        "target_url": url,
        "external_id": external_id,
        "detail": detail,
    }
