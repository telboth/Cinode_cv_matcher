from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document


def build_cinode_cv_docx(profile: dict[str, Any], resume: dict[str, Any] | None) -> bytes:
    document = Document()

    full_name = f"{profile.get('firstName', '')} {profile.get('lastName', '')}".strip() or "Consultant"
    title = str(profile.get("title") or "")
    location = str(profile.get("locationName") or "")
    email = str(profile.get("companyUserEmail") or "")

    document.add_heading(full_name, level=0)
    if title:
        document.add_paragraph(title)
    if location or email:
        contact = " | ".join([part for part in [location, email] if part])
        document.add_paragraph(contact)

    if resume and isinstance(resume.get("resume"), dict):
        resume_data = resume["resume"]
        blocks = resume_data.get("blocks")
        if isinstance(blocks, list):
            _append_summary_block(document, blocks)
            _append_skills_block(document, blocks)
            _append_other_blocks(document, blocks)

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _append_summary_block(document: Document, blocks: list[Any]) -> None:
    for block in blocks:
        if not isinstance(block, dict):
            continue
        if block.get("blockType") != 9:
            continue
        title = str(block.get("title") or "Sammendrag")
        description = str(block.get("description") or "").strip()
        if description:
            document.add_heading(title, level=1)
            document.add_paragraph(description)
        return


def _append_skills_block(document: Document, blocks: list[Any]) -> None:
    for block in blocks:
        if not isinstance(block, dict):
            continue
        if block.get("blockType") != 12:
            continue

        data = block.get("data")
        if not isinstance(data, list):
            continue

        skills: list[str] = []
        for category in data:
            if not isinstance(category, dict):
                continue
            entries = category.get("skills")
            if not isinstance(entries, list):
                continue
            for entry in entries:
                if not isinstance(entry, dict):
                    continue
                name = str(entry.get("name") or "").strip()
                if name:
                    skills.append(name)

        if skills:
            document.add_heading("Kompetanse", level=1)
            for skill in list(dict.fromkeys(skills))[:80]:
                document.add_paragraph(skill, style="List Bullet")
        return


def _append_other_blocks(document: Document, blocks: list[Any]) -> None:
    document.add_heading("CV-data", level=1)
    count = 0
    for block in blocks:
        if not isinstance(block, dict):
            continue
        if block.get("blockType") in {9, 12}:
            continue
        texts = _collect_candidate_texts(block)
        if not texts:
            continue
        count += 1
        heading = str(block.get("heading") or block.get("friendlyBlockName") or f"Seksjon {count}")
        document.add_heading(heading, level=2)
        for line in texts[:15]:
            document.add_paragraph(line)
        if count >= 4:
            break


def _collect_candidate_texts(node: Any) -> list[str]:
    keys = {"title", "description", "name", "subHeading", "employer", "assignment", "role", "text"}
    found: list[str] = []

    def walk(value: Any, depth: int = 0) -> None:
        if depth > 6:
            return
        if isinstance(value, dict):
            for key, item in value.items():
                if key in keys and isinstance(item, str):
                    text = item.strip()
                    if text:
                        found.append(text)
                walk(item, depth + 1)
        elif isinstance(value, list):
            for item in value:
                walk(item, depth + 1)

    walk(node)
    deduped = list(dict.fromkeys(found))
    return deduped
